"""Real-LLM API acceptance tests.

Run from the repository root:
    python -m unittest ai_core.tests.api_acceptance.test_real_llm_api_acceptance

These tests call FastAPI endpoints through TestClient and intentionally require
real LLM credentials from .env or the process environment.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

try:
    from dotenv import load_dotenv  # type: ignore
except Exception:  # pragma: no cover - optional local convenience dependency
    def load_dotenv(*_args, **_kwargs):
        return False

ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "ai_core" / "src"))
sys.path.insert(0, str(ROOT / "ai_core"))

load_dotenv(ROOT / ".env")

try:
    from docx import Document  # type: ignore
    from fastapi.testclient import TestClient  # type: ignore
    from openpyxl import Workbook, load_workbook  # type: ignore

    from main import app

    _IMPORT_ERROR = None
except Exception as exc:  # pragma: no cover - environment guard
    Document = None  # type: ignore
    TestClient = None  # type: ignore
    Workbook = None  # type: ignore
    load_workbook = None  # type: ignore
    app = None  # type: ignore
    _IMPORT_ERROR = exc


def _has_real_llm_config() -> bool:
    provider = os.getenv("LLM_PROVIDER", "zhipu").lower()
    if provider == "zhipu":
        return bool(os.getenv("ZHIPU_API_KEY"))
    return bool(os.getenv("OPENAI_API_KEY"))


def _write_doc_extract_source(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("项目名称：智慧办公自动填报平台")
    doc.add_paragraph("负责人：张三")
    doc.add_paragraph("预算：128万元")
    doc.add_paragraph("截止日期：2026年6月30日")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "项目状态"
    table.cell(1, 1).text = "验收中"
    doc.save(path)


def _write_template(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "空气质量"
    ws.append(["城市", "日期", "AQI", "PM2.5"])
    wb.save(path)


def _write_source_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "source"
    ws.append(["城市", "日期", "AQI", "PM2.5"])
    ws.append(["北京市", "2026-06-01", 90, 30])
    ws.append(["北京市", "2026-06-02", 130, 40])
    ws.append(["上海市", "2026-06-03", 180, 70])
    ws.append(["北京市", "2026-06-04", 160, 55])
    ws.append(["北京市", "2026-06-05", 120, 35])
    wb.save(path)


def _write_source_txt(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "北京市 2026-06-02 AQI 130 PM2.5 40。",
                "北京市 2026-06-04 AQI 160 PM2.5 55，是本周较高记录。",
                "上海市 2026-06-03 AQI 180 PM2.5 70，不属于本次城市筛选。",
            ]
        ),
        encoding="utf-8",
    )


@unittest.skipIf(_IMPORT_ERROR is not None, f"API acceptance dependencies unavailable: {_IMPORT_ERROR}")
@unittest.skipUnless(_has_real_llm_config(), "Real LLM API key is not configured in .env/environment.")
class RealLlmApiAcceptanceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(app)
        self.tmp = tempfile.TemporaryDirectory()
        self.work = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_doc_extract_upload_returns_metadata_and_values(self) -> None:
        source = self.work / "project.docx"
        _write_doc_extract_source(source)

        with source.open("rb") as file_obj:
            response = self.client.post(
                "/doc-extract/upload",
                data={"fields": "项目名称,负责人,预算,截止日期"},
                files={"file": ("project.docx", file_obj, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(payload["status"], "success")
        data = payload["extracted_data"]
        for field in ("项目名称", "负责人", "预算", "截止日期"):
            self.assertIn(field, data)
            self.assertNotIn(data[field], (None, "", "未找到"))
        self.assertIn("_meta", data)
        self.assertIn("confidence", data["_meta"])
        self.assertIn("validation", data["_meta"])

    def test_table_fill_upload_returns_filled_file_and_report(self) -> None:
        template = self.work / "template.xlsx"
        source_xlsx = self.work / "source.xlsx"
        source_txt = self.work / "source.txt"
        _write_template(template)
        _write_source_xlsx(source_xlsx)
        _write_source_txt(source_txt)
        report_dir = ROOT / "reports" / "table_fill"
        before_reports = set(report_dir.glob("*-fill-run-report.json")) if report_dir.exists() else set()

        request_text = "请筛选北京市2026年6月1日至2026年6月7日的数据，AQI大于100，按AQI降序取前2，只填写城市、日期、AQI和PM2.5。"
        with template.open("rb") as template_obj, source_xlsx.open("rb") as xlsx_obj, source_txt.open("rb") as txt_obj:
            response = self.client.post(
                "/table-fill/upload",
                data={"user_request": request_text},
                files=[
                    ("template", ("template.xlsx", template_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                    ("documents", ("source.xlsx", xlsx_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                    ("documents", ("source.txt", txt_obj, "text/plain")),
                ],
            )

        self.assertEqual(response.status_code, 200, response.text[:500])
        after_reports = set(report_dir.glob("*-fill-run-report.json")) if report_dir.exists() else set()
        new_reports = after_reports - before_reports
        self.assertTrue(new_reports, "table filling did not create a fill-run report")
        report_path = max(new_reports, key=lambda path: path.stat().st_mtime)
        self.assertTrue(report_path.exists(), f"report not found: {report_path}")

        output = self.work / "filled.xlsx"
        output.write_bytes(response.content)
        wb = load_workbook(output, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        filled_rows = [row for row in rows if any(value not in (None, "") for value in row)]
        self.assertGreaterEqual(len(filled_rows), 1)
        self.assertLessEqual(len(filled_rows), 2)
        self.assertTrue(all(row[0] == "北京市" for row in filled_rows))
        self.assertTrue(all(float(row[2]) > 100 for row in filled_rows if row[2] is not None))

        report = json.loads(report_path.read_text(encoding="utf-8"))
        self.assertGreaterEqual(report["input_summary"]["source_document_count"], 2)
        self.assertGreater(report["candidate_stats"]["merged_candidate_count"], 0)
        self.assertIn(report["verification_status"], {"pass", "warning"})
        self.assertIn("rag_result", report)
        self.assertIn("task_constraints", report)
        self.assertGreater(report["written_cell_count"], 0)


if __name__ == "__main__":
    unittest.main()
