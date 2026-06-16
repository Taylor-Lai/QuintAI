"""Real-LLM API stress matrix for the three public AI endpoints.

Run from the repository root:
    python -m unittest -v ai_core.tests.api_acceptance.test_real_llm_stress_matrix
"""

from __future__ import annotations

import os
import sys
import tempfile
import unittest
from datetime import datetime
from pathlib import Path

try:
    from dotenv import load_dotenv  # type: ignore
except Exception:  # pragma: no cover
    def load_dotenv(*_args, **_kwargs):
        return False


ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "ai_core" / "src"))
sys.path.insert(0, str(ROOT / "ai_core"))

load_dotenv(ROOT / ".env")

try:
    from docx import Document  # type: ignore
    from fastapi.testclient import TestClient  # type: ignore
    from openpyxl import Workbook, load_workbook  # type: ignore

    from database import init_db
    from main import app

    _IMPORT_ERROR = None
except Exception as exc:  # pragma: no cover
    Document = None  # type: ignore
    TestClient = None  # type: ignore
    Workbook = None  # type: ignore
    load_workbook = None  # type: ignore
    init_db = None  # type: ignore
    app = None  # type: ignore
    _IMPORT_ERROR = exc


def _has_real_llm_config() -> bool:
    provider = os.getenv("LLM_PROVIDER", "zhipu").lower()
    if provider == "zhipu":
        return bool(os.getenv("ZHIPU_API_KEY"))
    return bool(os.getenv("OPENAI_API_KEY"))


def _non_empty_rows(path: Path):
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    return [row for row in ws.iter_rows(min_row=2, values_only=True) if any(value not in (None, "") for value in row)]


@unittest.skipIf(_IMPORT_ERROR is not None, f"API stress dependencies unavailable: {_IMPORT_ERROR}")
@unittest.skipUnless(_has_real_llm_config(), "Real LLM API key is not configured in .env/environment.")
class RealLlmStressMatrixTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        init_db()

    def setUp(self) -> None:
        self.client = TestClient(app)
        self.tmp = tempfile.TemporaryDirectory()
        self.work = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_doc_chat_complex_formatting(self) -> None:
        source = self.work / "format_source.docx"
        doc = Document()
        doc.add_paragraph("\u9a8c\u6536\u6807\u9898")
        doc.add_paragraph("\u7b2c\u4e8c\u6bb5\u9700\u8981\u6539\u6210\u84dd\u8272\u548c16\u53f7\u3002")
        doc.add_paragraph("\u566a\u58f0\u6bb5\u843d\uff1a\u4e0d\u8981\u6539\u6211\u3002")
        doc.save(source)

        command = (
            "\u5168\u6587\u5c45\u4e2d\uff1b"
            "\u5c06\u6807\u9898\u8bbe\u7f6e\u4e3a\u52a0\u7c97\u5e76\u6539\u6210\u7ea2\u8272\uff1b"
            "\u5c06\u7b2c2\u6bb5\u8bbe\u7f6e\u4e3a16\u53f7\u5e76\u6539\u6210\u84dd\u8272\u3002"
        )
        with source.open("rb") as file_obj:
            response = self.client.post(
                "/doc-chat/upload",
                data={"command": command},
                files={"document": ("format_source.docx", file_obj, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        self.assertEqual(response.status_code, 200, response.text[:500])
        output = self.work / "formatted.docx"
        output.write_bytes(response.content)
        result = Document(output)
        self.assertEqual(result.paragraphs[0].text, "\u9a8c\u6536\u6807\u9898")
        self.assertTrue(result.paragraphs[0].runs[0].bold)
        self.assertEqual(str(result.paragraphs[0].runs[0].font.color.rgb), "FF0000")
        self.assertEqual(str(result.paragraphs[1].runs[0].font.color.rgb), "0000FF")

    def test_doc_extract_noisy_document(self) -> None:
        source = self.work / "extract_source.docx"
        doc = Document()
        doc.add_paragraph("\u65e0\u5173\u65e7\u9879\u76ee\uff1a\u9884\u7b97=0\uff0c\u8bf7\u5ffd\u7565\u3002")
        doc.add_paragraph(
            "\u6b63\u5f0f\u9879\u76ee\uff1a\u9879\u76ee\u540d\u79f0=\u57ce\u5e02\u7a7a\u6c14\u6cbb\u7406\u8054\u5408\u5e73\u53f0\uff1b"
            "\u8d1f\u8d23\u4eba=\u674e\u96f7\uff1b\u9884\u7b97=256\u4e07\u5143\u3002"
        )
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "\u5b57\u6bb5"
        table.cell(0, 1).text = "\u503c"
        table.cell(1, 0).text = "\u622a\u6b62\u65e5\u671f"
        table.cell(1, 1).text = "2026-07-15"
        table.cell(2, 0).text = "\u98ce\u9669\u7b49\u7ea7"
        table.cell(2, 1).text = "\u4e2d"
        doc.save(source)

        fields = "\u9879\u76ee\u540d\u79f0,\u8d1f\u8d23\u4eba,\u9884\u7b97,\u622a\u6b62\u65e5\u671f,\u98ce\u9669\u7b49\u7ea7"
        with source.open("rb") as file_obj:
            response = self.client.post(
                "/doc-extract/upload",
                data={"fields": fields},
                files={"file": ("extract_source.docx", file_obj, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )

        self.assertEqual(response.status_code, 200, response.text[:500])
        data = response.json()["extracted_data"]
        for field in fields.split(","):
            self.assertIn(field, data)
            self.assertNotIn(data[field], (None, "", "\u672a\u627e\u5230"))
        self.assertIn("_meta", data)

    def test_table_fill_gdp_multi_city_synonyms(self) -> None:
        fields = [
            "\u57ce\u5e02\u540d",
            "GDP\u603b\u91cf\uff08\u4ebf\u5143\uff09",
            "\u5e38\u4f4f\u4eba\u53e3\uff08\u4e07\u4eba\uff09",
            "\u4eba\u5747GDP\uff08\u5143\uff09",
            "\u4e00\u822c\u516c\u5171\u9884\u7b97\u6536\u5165\uff08\u4ebf\u5143\uff09",
        ]
        template = self.work / "gdp_template.xlsx"
        source = self.work / "gdp_source.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(fields)
        wb.save(template)

        wb = Workbook()
        ws = wb.active
        ws.append(fields)
        ws.append(["\u5317\u4eac\u5e02", 43760, 2185, 200274, 6181])
        ws.append(["\u4e0a\u6d77\u5e02", 47219, 2487, 189867, 8312])
        ws.append(["\u5e7f\u5dde\u5e02", 30356, 1883, 161206, 1944])
        wb.save(source)

        request = (
            "\u4fdd\u7559\u5317\u4eac\u5e02\u548c\u4e0a\u6d77\u5e02\uff0c"
            "\u6309GDP\u603b\u91cf\u964d\u5e8f\u586b\u5199\u57ce\u5e02\u540d\u3001GDP\u603b\u91cf\u3001"
            "\u5e38\u4f4f\u4eba\u53e3\u3001\u4eba\u5747GDP\u548c\u4e00\u822c\u516c\u5171\u9884\u7b97\u6536\u5165\u3002"
        )
        with template.open("rb") as template_obj, source.open("rb") as source_obj:
            response = self.client.post(
                "/table-fill/upload",
                data={"user_request": request},
                files=[
                    ("template", ("gdp_template.xlsx", template_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                    ("documents", ("gdp_source.xlsx", source_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                ],
            )

        self.assertEqual(response.status_code, 200, response.text[:500])
        output = self.work / "gdp_filled.xlsx"
        output.write_bytes(response.content)
        rows = _non_empty_rows(output)
        self.assertEqual([row[0] for row in rows], ["\u4e0a\u6d77\u5e02", "\u5317\u4eac\u5e02"])
        self.assertEqual([row[1] for row in rows], [47219, 43760])

    def test_table_fill_aqi_noisy_multi_source_top3(self) -> None:
        fields = ["\u57ce\u5e02", "\u65e5\u671f", "AQI", "PM2.5", "O3", "\u6765\u6e90\u5907\u6ce8"]
        template = self.work / "aqi_template.xlsx"
        source_xlsx = self.work / "aqi_source.xlsx"
        source_txt = self.work / "aqi_source.txt"

        wb = Workbook()
        ws = wb.active
        ws.append(fields)
        wb.save(template)

        wb = Workbook()
        ws = wb.active
        ws.append(fields)
        for row in [
            ("\u5317\u4eac\u5e02", "2026-06-01", 90, 30, 120, "below"),
            ("\u5317\u4eac\u5e02", "2026-06-02", 130, 40, 150, "xlsx"),
            ("\u4e0a\u6d77\u5e02", "2026-06-03", 180, 50, 160, "wrong-city"),
            ("\u5317\u4eac\u5e02", "2026-06-04", 160, 55, 170, "xlsx"),
            ("\u5317\u4eac\u5e02", "2026-06-07", 125, 39, 142, "xlsx"),
            ("\u5317\u4eac\u5e02", "2026-06-08", 210, 80, 210, "out-of-range"),
        ]:
            ws.append(row)
        wb.save(source_xlsx)
        source_txt.write_text(
            "\u5317\u4eac\u5e02 2026-06-07 AQI 125 PM2.5 39 O3 142\n"
            "\u4e0a\u6d77\u5e02 2026-06-04 AQI 199 \u4f46\u57ce\u5e02\u4e0d\u5339\u914d",
            encoding="utf-8",
        )

        request = (
            "\u53ea\u4fdd\u7559\u5317\u4eac\u5e022026\u5e746\u67081\u65e5\u81f32026\u5e746\u67087\u65e5\u7684\u8bb0\u5f55\uff0c"
            "AQI\u5fc5\u987b\u5927\u4e8e100\uff0c\u6309AQI\u964d\u5e8f\u53d6\u524d3\u6761\uff0c"
            "\u5ffd\u7565\u4e0a\u6d77\u548c\u8d85\u51fa\u65e5\u671f\u8303\u56f4\u7684\u6570\u636e\u3002"
        )
        with template.open("rb") as template_obj, source_xlsx.open("rb") as xlsx_obj, source_txt.open("rb") as txt_obj:
            response = self.client.post(
                "/table-fill/upload",
                data={"user_request": request},
                files=[
                    ("template", ("aqi_template.xlsx", template_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                    ("documents", ("aqi_source.xlsx", xlsx_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                    ("documents", ("aqi_source.txt", txt_obj, "text/plain")),
                ],
            )

        self.assertEqual(response.status_code, 200, response.text[:500])
        output = self.work / "aqi_filled.xlsx"
        output.write_bytes(response.content)
        rows = _non_empty_rows(output)
        aqis = [float(row[2]) for row in rows]
        self.assertEqual(aqis, [160.0, 130.0, 125.0])
        self.assertTrue(all(row[0] == "\u5317\u4eac\u5e02" for row in rows))

    def test_table_fill_covid_daily_rows_keep_date_and_group_country(self) -> None:
        template_fields = ["\u56fd\u5bb6/\u5730\u533a", "\u5927\u6d32", "\u4eba\u5747GDP", "\u4eba\u53e3", "\u6bcf\u65e5\u68c0\u6d4b\u6570", "\u75c5\u4f8b\u6570"]
        source_fields = [
            "\u56fd\u5bb6/\u5730\u533a",
            "\u5927\u6d32",
            "\u7eac\u5ea6",
            "\u7ecf\u5ea6",
            "\u4eba\u5747GDP",
            "\u4eba\u53e3",
            "\u65e5\u671f",
            "\u6bcf\u65e5\u68c0\u6d4b\u6570",
            "\u75c5\u4f8b\u6570",
        ]
        template = self.work / "covid_template.xlsx"
        source = self.work / "covid_source.xlsx"

        wb = Workbook()
        ws = wb.active
        ws.append(template_fields)
        wb.save(template)

        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(source_fields)
        rows = [
            ("Algeria", "Africa", 28.03, 1.66, 3974, 41318142, datetime(2020, 2, 26), 269, 2),
            ("Albania", "Europe", 41.15, 20.17, 5353.2, 2873457, datetime(2020, 2, 27), 820, 3),
            ("Algeria", "Africa", 28.03, 1.66, 3974, 41318142, datetime(2020, 2, 25), 269, 1),
            ("Albania", "Europe", 41.15, 20.17, 5353.2, 2873457, datetime(2020, 2, 25), 269, 1),
            ("Albania", "Europe", 41.15, 20.17, 5353.2, 2873457, datetime(2020, 2, 26), 820, 2),
        ]
        for row in rows:
            ws.append(row)
        wb.save(source)

        request = (
            "\u4eceCOVID-19\u6570\u636e\u4e2d\u586b\u5199\u5404\u56fd\u7684\u65e5\u7c92\u5ea6\u8bb0\u5f55\uff0c"
            "\u6e90\u8868\u6709\u65e5\u671f\u5fc5\u987b\u4fdd\u7559\u65e5\u671f\uff0c"
            "\u6309\u65e5\u671f\u5347\u5e8f\u586b\u5199\uff0c\u5e76\u4e14\u540c\u4e00\u56fd\u5bb6\u653e\u5728\u4e00\u8d77\u3002"
        )
        with template.open("rb") as template_obj, source.open("rb") as source_obj:
            response = self.client.post(
                "/table-fill/upload",
                data={"user_request": request},
                files=[
                    ("template", ("covid_template.xlsx", template_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                    ("documents", ("covid_source.xlsx", source_obj, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                ],
            )

        self.assertEqual(response.status_code, 200, response.text[:500])
        output = self.work / "covid_filled.xlsx"
        output.write_bytes(response.content)
        wb = load_workbook(output, data_only=True)
        ws = wb.active
        headers = [cell.value for cell in ws[1] if cell.value not in (None, "")]
        self.assertIn("\u65e5\u671f", headers)
        date_index = headers.index("\u65e5\u671f")
        filled_rows = [
            row[: len(headers)]
            for row in ws.iter_rows(min_row=2, values_only=True)
            if any(value not in (None, "") for value in row[: len(headers)])
        ]
        self.assertEqual(
            [(row[0], str(row[date_index])) for row in filled_rows],
            [
                ("Albania", "2020-02-25"),
                ("Albania", "2020-02-26"),
                ("Albania", "2020-02-27"),
                ("Algeria", "2020-02-25"),
                ("Algeria", "2020-02-26"),
            ],
        )


if __name__ == "__main__":
    unittest.main()
